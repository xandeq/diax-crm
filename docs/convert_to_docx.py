import os
import sys

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


def hex_to_rgb(hex_color):
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def convert_html_to_docx(html_path, docx_path):
    if not os.path.exists(html_path):
        print(f"Error: File {html_path} not found.")
        return

    with open(html_path, 'r', encoding='utf-8') as f:
        html_content = f.read()

    soup = BeautifulSoup(html_content, 'html.parser')
    doc = Document()

    # Set margins (A4 default is usually fine, but let's be safe)
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Header ---
    # Attempt to replicate simple text header
    header_div = soup.find('div', class_='modern-header')
    if header_div:
        header = section.header
        header_paragraph = header.paragraphs[0]
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Extract text from header (simplified)
        brand_name = header_div.find(class_='brand-name')
        brand_sub = header_div.find(class_='brand-sub')

        if brand_name:
            run = header_paragraph.add_run(brand_name.get_text().strip() + "\n")
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(15, 106, 66) # Match the green color roughly

        if brand_sub:
            run = header_paragraph.add_run(brand_sub.get_text().strip())
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(100, 100, 100)

    # --- Footer ---
    footer_div = soup.find('div', class_='doc-footer')
    if footer_div:
        footer = section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Just get all text from footer
        footer_text = footer_div.get_text(separator=' | ', strip=True)
        run = footer_paragraph.add_run(footer_text)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)

    # --- Content ---
    content_wrapper = soup.find('div', class_='content-wrapper')
    if not content_wrapper:
        print("Warning: .content-wrapper not found, using body")
        content_wrapper = soup.body

    # Process elements
    # We look for specific classes we generated

    # Title
    doc_title = content_wrapper.find(class_='doc-title')
    if doc_title:
        p = doc.add_paragraph()
        run = p.add_run(doc_title.get_text().strip())
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(15, 23, 42) # var(--primary)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph("") # Spacer

    # Iterate through other elements
    # We need to iterate children to keep order
    for element in content_wrapper.children:
        if element.name == 'div':
            if 'section' in element.get('class', []):
                h2 = element.find('h2')
                if h2:
                    heading = doc.add_heading(h2.get_text().strip(), level=1)
                    # Style the heading if needed, but default is usually okay
            elif 'topline' in element.get('class', []):
                pass # Already handled title separately, or ignore
            elif 'kv' in element.get('class', []):
                # Key-Value pairs
                p = doc.add_paragraph()
                text = element.get_text(separator=': ', strip=True)
                p.add_run(text)

        elif element.name == 'p':
            text = element.get_text().strip()
            if text:
                p = doc.add_paragraph(text)
                if 'small' in element.get('class', []):
                    p.style.font.size = Pt(9)

        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text().strip(), style='List Bullet')

        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text().strip(), style='List Number')

        elif element.name == 'hr':
            doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(docx_path)
    print(f"Successfully created {docx_path}")

if __name__ == "__main__":
    base_dir = os.path.dirname(os.path.abspath(__file__))
    input_html = os.path.join(base_dir, 'briefing.html')
    output_docx = os.path.join(base_dir, 'briefing.docx')

    if len(sys.argv) > 1:
        input_html = sys.argv[1]
    if len(sys.argv) > 2:
        output_docx = sys.argv[2]

    convert_html_to_docx(input_html, output_docx)
